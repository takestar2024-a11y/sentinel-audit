# -*- coding: utf-8 -*-
"""
サイバードッグ診断 - Word報告書ジェネレーター
scanner.full_scan() の結果 dict から日本語の .docx 報告書を生成する。

2モード:
  mode="quick" … クイック診断版（無料・Webからダウンロードする簡易サマリー）
  mode="full"  … 本診断報告書（納品用。改善アクション・手法・免責を詳述）

使い方（CLI）:
  python report.py example.com            # quick版を reports/ に生成
  python report.py example.com --full     # full版（納品用）を生成
"""
import io
import os
import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import scanner

# 深刻度の表示・色
SEV = {
    "c": {"label": "重大", "rgb": RGBColor(0xD1, 0x34, 0x38), "cell": "F8D7D7"},
    "w": {"label": "要改善", "rgb": RGBColor(0xB9, 0x7A, 0x0C), "cell": "FCEFD2"},
    "s": {"label": "良好", "rgb": RGBColor(0x1E, 0x8E, 0x6A), "cell": "D6F0E6"},
}

# 深刻度・項目に応じた改善ガイド（full版で使用）。タイトル部分一致で引く。
REMEDIATION = {
    "証明書の有効期限": "証明書の自動更新（ACME / Let's Encrypt 等）を設定し、失効前に確実に更新される運用にする。",
    "暗号化プロトコル": "サーバ設定で TLS 1.2 / 1.3 のみを許可し、TLS 1.0 / 1.1 を無効化する。",
    "旧プロトコルの無効化": "Webサーバ／ロードバランサで TLS 1.0・1.1 を明示的に無効化する。",
    "HTTPSリダイレクト": "HTTP(80番)へのアクセスを 301 で HTTPS へ常時リダイレクトする設定を追加する。",
    "HSTS": "レスポンスに Strict-Transport-Security ヘッダー（max-age=31536000 以上）を付与する。",
    "Content-Security-Policy": "コンテンツ読み込み元を制限する Content-Security-Policy を定義し、段階的に厳格化する。",
    "X-Frame-Options": "X-Frame-Options: DENY（または SAMEORIGIN）を付与し、クリックジャッキングを防止する。",
    "X-Content-Type-Options": "X-Content-Type-Options: nosniff を付与し、MIMEタイプ推測を防止する。",
    "Referrer-Policy": "Referrer-Policy（strict-origin-when-cross-origin 等）を設定し、リファラ漏えいを抑える。",
    "Permissions-Policy": "利用しないブラウザ機能を Permissions-Policy で明示的に無効化する。",
    "サーバ情報の露出": "Server ヘッダーからバージョン番号を除去し、ソフトウェア情報の露出を抑える。",
    "SPF": "自社の正規送信元のみを許可する SPF レコードを設定し、終端を -all（厳格）にする。",
    "DKIM": "メール配信基盤で DKIM 署名を有効化し、公開鍵を DNS に公開する。",
    "DMARC": "DMARC レコードを設定し、監視（p=none）から段階的に p=quarantine / reject へ引き上げる。",
    "類似ドメイン登録": "検出された類似ドメインの用途を確認し、必要に応じて防衛的取得・監視・法的対応を検討する。",
}


def _grade(score):
    if score >= 75:
        return ("A", "良好", RGBColor(0x1E, 0x8E, 0x6A))
    if score >= 50:
        return ("B", "要改善", RGBColor(0xB9, 0x7A, 0x0C))
    return ("C", "危険", RGBColor(0xD1, 0x34, 0x38))


def _shade(cell, hex_color):
    """テーブルセルに背景色を付ける。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_font(doc):
    """既定フォントを日本語対応に。"""
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "Yu Gothic")


def _remediation_for(title):
    for key, adv in REMEDIATION.items():
        if key in title:
            return adv
    return "設定内容を確認し、ベストプラクティスに沿って是正する。"


def build_document(report, mode="quick"):
    """scanレポート dict から Document を組み立てて返す。"""
    full = (mode == "full")
    doc = Document()
    _set_font(doc)

    # 余白
    for s in doc.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

    g, gtext, gcolor = _grade(report["overall"])

    # ---- タイトル ----
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = brand.add_run("サイバードッグ診断")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x22, 0x9E, 0xB8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Webセキュリティ診断報告書" if full else "Webセキュリティ クイック診断結果")
    tr.bold = True; tr.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("本診断報告書（詳細版）" if full else "無償クイック診断版（ダイジェスト）")
    sr.font.size = Pt(10.5); sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    credit = doc.add_paragraph()
    credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = credit.add_run("診断監修：ドッカー先生 🐾")
    cr.italic = True; cr.font.size = Pt(9.5); cr.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)

    doc.add_paragraph()

    # ---- 概要テーブル ----
    meta = doc.add_table(rows=0, cols=2)
    meta.style = "Light List Accent 1"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    def row(k, v, color=None, bold=False):
        cells = meta.add_row().cells
        cells[0].paragraphs[0].add_run(k).bold = True
        rr = cells[1].paragraphs[0].add_run(v)
        if bold: rr.bold = True
        if color is not None: rr.font.color.rgb = color
        rr.font.size = Pt(11)
    row("診断対象ドメイン", report["domain"], bold=True)
    row("診断日時", report.get("scannedAt", datetime.datetime.now().strftime("%Y-%m-%d %H:%M")))
    row("総合スコア", f"{report['overall']} / 100 点")
    row("総合評価", f"{g}（{gtext}）", color=gcolor, bold=True)
    c = report["counts"]
    row("検出内訳", f"重大 {c['c']}件 ／ 要改善 {c['w']}件 ／ 良好 {c['s']}件")

    doc.add_paragraph()

    # ---- 総評 ----
    doc.add_heading("総評", level=1)
    if c["c"] > 0:
        summary = (f"外部から観測した結果、重大な指摘が {c['c']} 件、要改善が {c['w']} 件検出されました。"
                   "特になりすまし・暗号化に関わる項目は、被害が発生する前の早急な是正を推奨します。")
    elif c["w"] > 0:
        summary = (f"致命的な問題は検出されませんでしたが、要改善の項目が {c['w']} 件あります。"
                   "攻撃の足がかりを減らすため、順次対応することを推奨します。")
    else:
        summary = ("主要な項目は良好な状態です。設定の劣化や新たな偽装ドメインの出現を防ぐため、"
                   "定期的な再診断を推奨します。")
    doc.add_paragraph(summary)

    # ---- 領域別詳細 ----
    doc.add_heading("領域別 診断結果", level=1)
    for area in report["areaResults"]:
        h = doc.add_heading(level=2)
        hr = h.add_run(f"{area['name']}（{area['score']} / 100）")
        ag, _, agc = _grade(area["score"])
        hr.font.color.rgb = agc

        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, t in enumerate(("深刻度", "項目", "内容")):
            run = hdr[i].paragraphs[0].add_run(t); run.bold = True
            _shade(hdr[i], "E7ECF2")
        for f in area["findings"]:
            cells = tbl.add_row().cells
            sev = SEV[f["sev"]]
            sr = cells[0].paragraphs[0].add_run(sev["label"])
            sr.bold = True; sr.font.color.rgb = sev["rgb"]
            _shade(cells[0], sev["cell"])
            cells[1].paragraphs[0].add_run(f["title"])
            cells[2].paragraphs[0].add_run(f["desc"])
        # 列幅
        for row_ in tbl.rows:
            row_.cells[0].width = Cm(2.2); row_.cells[1].width = Cm(4.5); row_.cells[2].width = Cm(9.5)
        doc.add_paragraph()

    # ---- full版：改善アクション（優先度順）----
    if full:
        doc.add_heading("改善アクション（優先度順）", level=1)
        actions = []
        for area in report["areaResults"]:
            for f in area["findings"]:
                if f["sev"] in ("c", "w"):
                    actions.append((f["sev"], area["name"], f["title"]))
        actions.sort(key=lambda x: 0 if x[0] == "c" else 1)
        if not actions:
            doc.add_paragraph("優先対応が必要な指摘は検出されませんでした。現状維持と定期監視を推奨します。")
        else:
            atbl = doc.add_table(rows=1, cols=4)
            atbl.style = "Table Grid"
            hdr = atbl.rows[0].cells
            for i, t in enumerate(("優先度", "領域", "項目", "推奨対応")):
                run = hdr[i].paragraphs[0].add_run(t); run.bold = True
                _shade(hdr[i], "E7ECF2")
            for sev, area_name, title in actions:
                cells = atbl.add_row().cells
                s = SEV[sev]
                pr = cells[0].paragraphs[0].add_run("最優先" if sev == "c" else "早期")
                pr.bold = True; pr.font.color.rgb = s["rgb"]; _shade(cells[0], s["cell"])
                cells[1].paragraphs[0].add_run(area_name)
                cells[2].paragraphs[0].add_run(title)
                cells[3].paragraphs[0].add_run(_remediation_for(title))
            for row_ in atbl.rows:
                row_.cells[0].width = Cm(1.8); row_.cells[1].width = Cm(3.2)
                row_.cells[2].width = Cm(4.0); row_.cells[3].width = Cm(7.2)
        doc.add_paragraph()

        doc.add_heading("診断範囲と手法", level=1)
        doc.add_paragraph(
            "本診断は、外部から観測可能な公開情報のみを用いた非侵入型の実測診断です。"
            "対象サーバーへの侵入・負荷はかけていません。以下の5領域を観測しました：")
        for name in ("SSL/TLS証明書（有効期限・発行元・プロトコル・旧TLSの受け入れ・HTTPS転送）",
                     "HTTPセキュリティヘッダー（HSTS / CSP / X-Frame-Options ほか）",
                     "DNSメール認証（SPF / DKIM / DMARC）",
                     "類似・偽装ドメインの登録・稼働状況"):
            doc.add_paragraph(name, style="List Bullet")

    # ---- CTA / 免責 ----
    doc.add_paragraph()
    if not full:
        cta = doc.add_heading("全項目の詳細診断について", level=1)
        doc.add_paragraph(
            "本書は無償クイック診断のダイジェストです。本診断では40項目以上をフルチェックし、"
            "優先順位付きの改善手順を含む詳細報告書を納品します。ご希望の際は下記までご連絡ください。")
        doc.add_paragraph("お問い合わせ（公式LINE）：https://lin.ee/EYNlG8m")
        doc.add_paragraph("メール：take.star2024@gmail.com")

    disc = doc.add_paragraph()
    dr = disc.add_run(
        "【免責】本報告書は観測時点の状態に基づく非侵入型診断の結果であり、ペネトレーションテスト"
        "（侵入テスト）は含みません。DKIMは代表的なセレクタのみを照会するため、未検出でも設定済みの"
        "場合があります。類似ドメインの検出は登録・稼働の有無を示すもので悪性を断定するものではありません。"
        "本書はすべての脆弱性の不存在を保証するものではありません。")
    dr.font.size = Pt(8.5); dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


def generate_bytes(report, mode="quick"):
    """docx を bytes で返す（Webダウンロード用）。"""
    doc = build_document(report, mode)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _out_dir():
    d = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
    os.makedirs(d, exist_ok=True)
    return d


def generate_file(domain, mode="quick"):
    """スキャンを実行し、reports/ に .docx を保存してパスを返す。"""
    report = scanner.full_scan(domain)
    doc = build_document(report, mode)
    suffix = "本診断報告書" if mode == "full" else "クイック診断"
    safe = domain.replace(":", "_").replace("/", "_")
    path = os.path.join(_out_dir(), f"{safe}_{suffix}.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    import sys
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    mode = "full" if "--full" in sys.argv else "quick"
    domain = args[0] if args else "example.com"
    ok, reason = scanner.is_public_domain(domain)
    if not ok:
        print(f"診断できません（{reason}）：{domain}")
        sys.exit(1)
    p = generate_file(domain, mode)
    print(f"生成しました（{mode}）：{p}")
