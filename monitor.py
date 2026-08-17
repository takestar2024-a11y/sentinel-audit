# -*- coding: utf-8 -*-
"""
SiteDoc AI - 月額保守監査モニター
「③ 月額保守監査」商品の技術的な裏付け。前回の診断結果と比較し、
新たに検出された問題・解消された問題・新規の類似ドメイン登録・
DNS応答の整合性（複数リゾルバ照合）をまとめ、月次保守レポート（Word）を生成する。

Webサーバー(server.py)はRender無料プランのステートレスな環境（アイドル時スリープ・
DB無し）のため、継続監視の状態はここでは持たない。report.py --full / excel_report.py
と同じく、オペレーターが月次でCLI実行する運用を前提とする。

使い方（CLI）:
  python monitor.py example.com              # 初回はベースライン作成、2回目以降は差分レポート
  python monitor.py example.com --report     # 差分レポートをWordで reports/ に保存
"""
import io
import os
import json
import datetime

import dns.resolver
import dns.exception
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import scanner

STATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "monitor_state")

# 監視に使う公開DNSリゾルバ（DNS侵害＝経路上でのDNS応答すり替えの検知に使う）
_PUBLIC_RESOLVERS = {
    "Google": "8.8.8.8",
    "Cloudflare": "1.1.1.1",
    "Quad9": "9.9.9.9",
}


# ============================================================
# 状態の保存・読込
# ============================================================
def _state_path(domain):
    safe = domain.replace(":", "_").replace("/", "_")
    return os.path.join(STATE_DIR, f"{safe}.json")


def load_previous(domain):
    path = _state_path(domain)
    if not os.path.exists(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def save_state(domain, report, dns_records, lookalikes):
    os.makedirs(STATE_DIR, exist_ok=True)
    state = {
        "domain": domain,
        "saved_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "report": report,
        "dns_records": dns_records,
        "lookalikes": lookalikes,
    }
    with open(_state_path(domain), "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)


# ============================================================
# ① スキャン結果の差分（今月 vs 前回）
# ============================================================
def diff_findings(prev_report, curr_report):
    """前回・今回の findings を (領域名, 項目) キーで比較する。"""

    def _index(report):
        idx = {}
        for area in report.get("areaResults", []):
            for f in area["findings"]:
                idx[(area["name"], f["title"])] = f
        return idx

    prev_idx = _index(prev_report)
    curr_idx = _index(curr_report)

    new_issues = []      # 前回無かった c/w が今回出現
    resolved_issues = []  # 前回あった c/w が今回は無い、または s に改善
    for key, f in curr_idx.items():
        if f["sev"] in ("c", "w") and key not in prev_idx:
            new_issues.append((key[0], f))
    for key, f in prev_idx.items():
        if f["sev"] in ("c", "w"):
            curr_f = curr_idx.get(key)
            if curr_f is None or curr_f["sev"] == "s":
                resolved_issues.append((key[0], f))

    return {
        "new_issues": new_issues,
        "resolved_issues": resolved_issues,
        "score_before": prev_report.get("overall", 0),
        "score_after": curr_report.get("overall", 0),
        "prev_scanned_at": prev_report.get("scannedAt", ""),
    }


# ============================================================
# ② 類似ドメインの新規登録・武装化（MX新規獲得）監視
# ============================================================
def _resolves_retry(name, attempts=2):
    """月次監視は正確性を優先するため、DNSの一時的な取りこぼしを減らすリトライを入れる。
    scanner.py本体（単発スキャンのレイテンシ最優先）には影響させない。"""
    for _ in range(attempts):
        if scanner._resolves(name):
            return True
    return False


def _mx_retry(name, attempts=2):
    for _ in range(attempts):
        if scanner._mx(name):
            return True
    return False


def check_lookalike_changes(domain, prev_state):
    """前回把握していた類似ドメイン一覧と比べ、新規登録・MX新規獲得を検知する。
    候補生成は scanner.py の実装を再利用する（二重管理を避けるため）。
    ただしDNS照会はリトライ付きのラッパーを使う（下記コメント参照）。"""
    cands = scanner._lookalike_candidates(domain)
    import concurrent.futures
    with concurrent.futures.ThreadPoolExecutor(max_workers=16) as ex:
        resolved = dict(zip(cands, ex.map(_resolves_retry, cands)))
    registered = [c for c, r in resolved.items() if r]
    with concurrent.futures.ThreadPoolExecutor(max_workers=16) as ex:
        mx_map = dict(zip(registered, ex.map(_mx_retry, registered)))

    current = {c: {"has_mail": mx_map.get(c, False)} for c in registered}

    prev_lookalikes = (prev_state or {}).get("lookalikes", {}) if prev_state else {}

    new_registrations = [c for c in current if c not in prev_lookalikes]
    newly_mail_capable = [
        c for c in current
        if c in prev_lookalikes and current[c]["has_mail"] and not prev_lookalikes[c].get("has_mail", False)
    ]

    return {
        "current": current,
        "new_registrations": new_registrations,
        "newly_mail_capable": newly_mail_capable,
        "total_registered": len(current),
    }


# ============================================================
# ③ DNS応答の整合性（複数リゾルバ照合）
# ============================================================
def _resolve_with(nameserver, name, rtype):
    try:
        r = dns.resolver.Resolver(configure=False)
        r.nameservers = [nameserver]
        r.timeout = 3.0
        r.lifetime = 3.0
        answers = r.resolve(name, rtype)
        return sorted(str(a).rstrip(".").lower() for a in answers)
    except (dns.resolver.NXDOMAIN, dns.resolver.NoAnswer,
            dns.resolver.NoNameservers, dns.exception.Timeout,
            dns.resolver.LifetimeTimeout, Exception):
        return []


def check_dns_integrity(domain):
    """複数の公開DNSに同じ質問をして、NSの回答が食い違っていないか確認する。
    NSの不一致はドメイン乗っ取り・経路上でのDNS応答すり替えの強い兆候。
    Aレコードの不一致はCDN等で正常に起こり得るため参考情報に留める。"""
    ns_by_resolver = {}
    a_by_resolver = {}
    for label, ip in _PUBLIC_RESOLVERS.items():
        ns_by_resolver[label] = tuple(_resolve_with(ip, domain, "NS"))
        a_by_resolver[label] = tuple(_resolve_with(ip, domain, "A"))

    ns_values = [v for v in ns_by_resolver.values() if v]
    ns_mismatch = len(set(ns_values)) > 1 if ns_values else False

    a_values = [v for v in a_by_resolver.values() if v]
    a_mismatch = len(set(a_values)) > 1 if a_values else False

    current_ns = ns_values[0] if ns_values and not ns_mismatch else None
    return {
        "ns_by_resolver": ns_by_resolver,
        "a_by_resolver": a_by_resolver,
        "ns_mismatch": ns_mismatch,
        "a_mismatch": a_mismatch,
        "current_ns": list(current_ns) if current_ns else None,
    }


def check_ns_change(domain, prev_state, dns_integrity_result):
    """前回保存したNSと、今回確認できたNSを比較する（乗っ取りの最重要シグナル）。"""
    prev_ns = (prev_state or {}).get("dns_records", {}).get("ns") if prev_state else None
    curr_ns = dns_integrity_result.get("current_ns")
    if prev_ns is None or curr_ns is None:
        return {"changed": False, "prev_ns": prev_ns, "curr_ns": curr_ns}
    return {"changed": sorted(prev_ns) != sorted(curr_ns), "prev_ns": prev_ns, "curr_ns": curr_ns}


# ============================================================
# 月次保守レポート（Word）
# ============================================================
def _set_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "Yu Gothic")


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_report_document(domain, diff, lookalike_result, dns_integrity_result, ns_change):
    doc = Document()
    _set_font(doc)
    for s in doc.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = brand.add_run("SiteDoc AI")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x22, 0x9E, 0xB8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("月次保守レポート")
    tr.bold = True; tr.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"{domain}｜{datetime.datetime.now().strftime('%Y年%m月%d日')} 確認分")
    sr.font.size = Pt(10.5); sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # ---- ①スコアの推移 ----
    doc.add_heading("① 前回からの変化", level=1)
    if diff["prev_scanned_at"]:
        p = doc.add_paragraph()
        p.add_run(f"前回確認: {diff['prev_scanned_at']}　→　今回: "
                   f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(10.5)
        p2 = doc.add_paragraph()
        delta = diff["score_after"] - diff["score_before"]
        sign = "+" if delta >= 0 else ""
        p2.add_run(f"総合スコア: {diff['score_before']} → {diff['score_after']} 点"
                    f"（{sign}{delta}）").bold = True

        if diff["new_issues"]:
            doc.add_heading("新たに検出された問題", level=2)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            for i, t in enumerate(("深刻度", "領域", "項目")):
                run = hdr[i].paragraphs[0].add_run(t); run.bold = True
                _shade(hdr[i], "E7ECF2")
            for area_name, f in diff["new_issues"]:
                cells = tbl.add_row().cells
                sev = {"c": ("重大", RGBColor(0xD1, 0x34, 0x38), "F8D7D7"),
                       "w": ("要改善", RGBColor(0xB9, 0x7A, 0x0C), "FCEFD2")}[f["sev"]]
                run = cells[0].paragraphs[0].add_run(sev[0]); run.bold = True; run.font.color.rgb = sev[1]
                _shade(cells[0], sev[2])
                cells[1].paragraphs[0].add_run(area_name)
                cells[2].paragraphs[0].add_run(f["title"])
        else:
            doc.add_paragraph("新たに検出された問題はありません。")

        if diff["resolved_issues"]:
            doc.add_heading("改善が確認された項目", level=2)
            for area_name, f in diff["resolved_issues"]:
                doc.add_paragraph(f"{area_name} — {f['title']}", style="List Bullet")
    else:
        doc.add_paragraph("今回が初回の記録です。次回以降、変化を比較して報告します。")

    doc.add_paragraph()

    # ---- ②類似ドメイン監視 ----
    doc.add_heading("② 類似ドメインの監視", level=1)
    doc.add_paragraph(f"現在、登録・稼働中の類似ドメイン: {lookalike_result['total_registered']}件")
    if lookalike_result["newly_mail_capable"]:
        p = doc.add_paragraph()
        run = p.add_run("🚨 メール送信能力を新たに獲得した類似ドメインがあります（最優先で確認してください）:")
        run.bold = True; run.font.color.rgb = RGBColor(0xD1, 0x34, 0x38)
        for d in lookalike_result["newly_mail_capable"]:
            doc.add_paragraph(d, style="List Bullet")
    if lookalike_result["new_registrations"]:
        doc.add_paragraph("今回新たに登録・稼働を確認した類似ドメイン:")
        for d in lookalike_result["new_registrations"][:10]:
            doc.add_paragraph(d, style="List Bullet")
    if not lookalike_result["newly_mail_capable"] and not lookalike_result["new_registrations"]:
        doc.add_paragraph("新規の登録・変化はありませんでした。")

    doc.add_paragraph()

    # ---- ③DNS整合性 ----
    doc.add_heading("③ DNS応答の整合性（DNS侵害の監視）", level=1)
    doc.add_paragraph(
        "複数の公開DNS（Google/Cloudflare/Quad9）に同じ質問をして、回答が"
        "食い違っていないかを確認しています。委任先(NS)の食い違いは、正しいURLでも"
        "偽サイトへ誘導される「DNS侵害」やドメイン乗っ取りの強い兆候です。"
    )
    if dns_integrity_result["ns_mismatch"]:
        p = doc.add_paragraph()
        run = p.add_run("🚨 委任先(NS)の回答がリゾルバ間で食い違っています。詳細な確認を推奨します。")
        run.bold = True; run.font.color.rgb = RGBColor(0xD1, 0x34, 0x38)
    elif ns_change["changed"]:
        p = doc.add_paragraph()
        run = p.add_run("⚠ 前回確認時からNS(委任先)が変更されています。心当たりのある変更か確認してください。")
        run.bold = True; run.font.color.rgb = RGBColor(0xB9, 0x7A, 0x0C)
    else:
        doc.add_paragraph("委任先(NS)の回答は各リゾルバで一致しており、異常は確認されませんでした。")

    doc.add_paragraph()
    disc = doc.add_paragraph()
    dr = disc.add_run(
        "【免責】本レポートは外部から観測可能な公開情報のみに基づく、非侵入型の定期確認結果です。"
        "すべての脅威の不存在を保証するものではありません。緊急性の高い指摘（🚨）を検出した場合は、"
        "登録者情報の確認やDNS設定の見直しなど、速やかな一次対応を推奨します。")
    dr.font.size = Pt(8.5); dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


# ============================================================
# CLI
# ============================================================
def _out_dir():
    d = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
    os.makedirs(d, exist_ok=True)
    return d


def run(domain, make_report=False):
    ok, reason = scanner.is_public_domain(domain)
    if not ok:
        print(f"診断できません（{reason}）：{domain}")
        return None

    prev_state = load_previous(domain)
    curr_report = scanner.full_scan(domain)
    dns_integrity_result = check_dns_integrity(domain)
    lookalike_result = check_lookalike_changes(domain, prev_state)
    ns_change = check_ns_change(domain, prev_state, dns_integrity_result)

    if prev_state is None:
        print(f"■ [{domain}] 初回記録を作成しました。総合スコア {curr_report['overall']}/100。"
              f"次回以降、変化を比較して報告します。")
        save_state(domain, curr_report,
                   {"ns": dns_integrity_result.get("current_ns")},
                   lookalike_result["current"])
        return None

    diff = diff_findings(prev_state["report"], curr_report)

    print(f"■ [{domain}] スコア {diff['score_before']} → {diff['score_after']}　"
          f"新規問題 {len(diff['new_issues'])}件　解消 {len(diff['resolved_issues'])}件")
    print(f"   類似ドメイン: 登録{lookalike_result['total_registered']}件"
          f"（新規{len(lookalike_result['new_registrations'])}件　"
          f"MX新規獲得{len(lookalike_result['newly_mail_capable'])}件）")
    if dns_integrity_result["ns_mismatch"] or ns_change["changed"]:
        print("   🚨 DNS整合性に注意: NSの不一致または変更を検出")
    else:
        print("   DNS整合性: 異常なし")

    if make_report:
        doc = build_report_document(domain, diff, lookalike_result, dns_integrity_result, ns_change)
        safe = domain.replace(":", "_").replace("/", "_")
        path = os.path.join(_out_dir(), f"{safe}_月次保守レポート.docx")
        doc.save(path)
        print(f"■ 月次保守レポートを生成しました：{path}")

    # 状態を更新（次回比較用）
    save_state(domain, curr_report,
               {"ns": dns_integrity_result.get("current_ns")},
               lookalike_result["current"])

    return diff


if __name__ == "__main__":
    import sys
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    make_report = "--report" in sys.argv
    domain = args[0] if args else "example.com"
    run(domain, make_report=make_report)
